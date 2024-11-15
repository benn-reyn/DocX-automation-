from docx import Document
from docxtpl import DocxTemplate
import pandas as pd

def read_word_file(file_path):
    doc = Document(file_path)
    data = {
        'prepared_for': 'Derrick Plumbing, Inc.',
        'effective_dates': '',
        'pollution_liability_dates': '',
        'as_of_date': '',
        'totalPremium': '',
        'company_name': 'Propel Insurance',
        'parsing_error': 'ERROR'
    }
    
    table_data = {}
    table_counter = 0
    
    paragraph_count = len(doc.paragraphs)
    
    for i in range(paragraph_count):
        para = doc.paragraphs[i]
        text = para.text.strip()
        print(f"Reading paragraph: {text}")  

        if not text:
            continue

        if "prepared for" in text.lower():
            if i + 1 < paragraph_count:
                data['prepared_for'] = doc.paragraphs[i + 1].text.strip()
        
        if "effective dates:" in text.lower():
            if i + 1 < paragraph_count:
                data['effective_dates'] = doc.paragraphs[i + 1].text.strip()
        
        if "non-Concurrent Pollution Liability effective" in text.lower():
            data['pollution_liability_dates'] = text.replace("Non-Concurrent Pollution Liability effective", "").strip()
        
        if text.lower().startswith("as of"):
            data['as_of_date'] = text.replace("as of", "").strip()
        
        if "premium summary" in text.lower():
            table0 = doc.tables[0]

            for i, row in enumerate(table0.rows):
                for j, cell in enumerate(row.cells):
                    placeholder = f"table0_cell_{i}_{j}"
                    data[placeholder] = cell.text
                    print(f"{placeholder}: {cell.text}")
                    
        if "estimated annual premium:" in text.lower():
            data['totalPremium'] = text.replace("Estimated Annual Premium:", "").strip()
            
        if "location 1:  2226 ridge road, batesburg leesville, sc" in text.lower():
            table1 = doc.tables[1]
            
            for i, row in enumerate(table1.rows):
                for j, cell in enumerate(row.cells):
                    placeholder = f"table1_cell_{i}_{j}"
                    data[placeholder] = cell.text
                    print(f"{placeholder}: {cell.text}")
                            

    if not data['prepared_for']:
        data['prepared_for'] = data['parsing_error']

    print(f"Parsed data: {data}")
    return data

def prepare_context(data):
    context = {key: value for key, value in data.items() if key.startswith("table0_") or key.startswith("table1_")}
    context.update({
        'prepared_for': data['prepared_for'],
        'effective_dates': data['effective_dates'],
        'pollution_liability_dates': data['pollution_liability_dates'],
        'as_of_date': data['as_of_date'],
        'totalPremium': data['totalPremium'],
        'company_name': data['company_name']
    })
    return context

def fill_template(template_path, context, output_path):
    doc = DocxTemplate(template_path)
    print(f"Context to fill the template: {context}") 
    doc.render(context)
    doc.save(output_path)

# Read the data from the input document
input_data = read_word_file('C:\\Users\\reyno\\OneDrive\\Desktop\\buis proj\\Old large POI template.docx')

# Prepare the context for the template
context = prepare_context(input_data)

# Fill the template with the context data
fill_template('C:\\Users\\reyno\\OneDrive\\Desktop\\buis proj\\Proposal of Insurance Large template.docx', context, 'C:\\Users\\reyno\\OneDrive\\Desktop\\buis proj\\generated_doc.docx')
